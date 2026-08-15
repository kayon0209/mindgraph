import io
import sys
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter
from reportlab.pdfgen import canvas

from application.structured_chunker import StructuredChunker
from infrastructure.parsers import default_parser_registry


class DocumentParserTests(unittest.TestCase):
    def test_markdown_and_txt(self):
        parsed = default_parser_registry.parse("# 差旅\n## 标准\n第一条 应及时报销".encode(), "policy.md")
        self.assertEqual([item.element_type for item in parsed.elements], ["heading", "heading", "numbered_clause"])
        self.assertEqual(parsed.elements[-1].heading_path, ["差旅", "标准"])
        self.assertEqual(default_parser_registry.parse("普通文本".encode(), "x.txt").elements[0].text, "普通文本")
        with self.assertRaises(ValueError): default_parser_registry.parse(b"\xff", "bad.txt")

    def test_text_pdf_and_scanned_detection(self):
        buffer = io.BytesIO(); pdf = canvas.Canvas(buffer); pdf.drawString(72, 720, "Travel expense policy requires valid invoice and approval materials."); pdf.save()
        parsed = default_parser_registry.parse(buffer.getvalue(), "policy.pdf")
        self.assertEqual(parsed.metadata["page_count"], 1); self.assertFalse(parsed.ocr_required_pages)
        blank = io.BytesIO(); writer = PdfWriter(); writer.add_blank_page(300, 300); writer.write(blank)
        scanned = default_parser_registry.parse(blank.getvalue(), "scan.pdf")
        self.assertEqual(scanned.ocr_required_pages, [1]); self.assertTrue(scanned.warnings)

    def test_docx_headings_lists_and_tables(self):
        document = Document(); document.add_heading("费用制度", level=1); document.add_paragraph("第一条 提交发票"); document.add_paragraph("材料", style="List Bullet")
        table = document.add_table(rows=2, cols=2); table.cell(0, 0).text = "项目"; table.cell(0, 1).text = "标准"; table.cell(1, 0).text = "住宿"; table.cell(1, 1).text = "500"
        buffer = io.BytesIO(); document.save(buffer)
        parsed = default_parser_registry.parse(buffer.getvalue(), "policy.docx")
        self.assertTrue(any(item.element_type == "heading" for item in parsed.elements)); self.assertTrue(any(item.element_type == "table" for item in parsed.elements))

    def test_xlsx_preserves_sheet_and_rows(self):
        workbook = Workbook(); sheet = workbook.active; sheet.title = "标准"; sheet.append(["城市", "金额"]); sheet.append(["上海", 500]); buffer = io.BytesIO(); workbook.save(buffer)
        parsed = default_parser_registry.parse(buffer.getvalue(), "policy.xlsx")
        self.assertEqual(parsed.elements[0].table_id, "sheet:标准"); self.assertIn("上海 | 500", parsed.elements[0].text)

    def test_corrupt_and_unsupported(self):
        for name in ("bad.pdf", "bad.docx", "bad.xlsx"):
            with self.assertRaises(ValueError): default_parser_registry.parse(b"broken", name)
        with self.assertRaises(ValueError): default_parser_registry.parse(b"x", "x.exe")

    def test_parent_child_chunking_preserves_metadata(self):
        parsed = default_parser_registry.parse(("# 制度\n第一条 " + "报销条件。" * 200).encode(), "long.md")
        chunks = StructuredChunker(child_size=120, parent_size=500, overlap=20).chunk(parsed)
        self.assertGreater(len(chunks), 1); self.assertTrue(all(item.parent_chunk_id and item.child_chunk_id for item in chunks)); self.assertTrue(chunks[0].heading_path)


if __name__ == "__main__": unittest.main()
