from __future__ import annotations

import hashlib
import re
from io import BytesIO

from docx import Document

from domain.models import ElementType, ParsedDocument, ParsedElement


class DOCXParser:
    name, version = "python-docx", "1.0.0"

    def supports(self, file_type: str) -> bool: return file_type.lower() == "docx"

    def parse(self, data: bytes, document_name: str) -> ParsedDocument:
        try: document = Document(BytesIO(data))
        except Exception as exc: raise ValueError("Corrupted DOCX") from exc
        elements: list[ParsedElement] = []
        heading_path: list[str] = []
        order = 0
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text: continue
            style = paragraph.style.name if paragraph.style else ""
            heading = re.match(r"Heading (\d+)", style, re.I)
            kind: ElementType
            if heading:
                level = int(heading.group(1)); heading_path = heading_path[:level - 1] + [text]; kind = "heading"
            elif style.lower().startswith("list"):
                kind = "list_item"
            elif re.match(r"^(第.+条|\d+(?:\.\d+)*[、.])", text):
                kind = "numbered_clause"
            else: kind = "paragraph"
            elements.append(ParsedElement(element_type=kind, text=text, order=order, heading_path=list(heading_path))); order += 1
        for table_index, table in enumerate(document.tables, 1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            text = "\n".join(" | ".join(row) for row in rows)
            elements.append(ParsedElement(element_type="table", text=text, order=order, heading_path=list(heading_path), table_id=f"table-{table_index}", table_rows=rows)); order += 1
        properties = document.core_properties
        return ParsedDocument(document_id=hashlib.sha256(document_name.encode()).hexdigest()[:16], document_name=document_name,
            file_type="docx", checksum=hashlib.sha256(data).hexdigest(), parser_name=self.name, parser_version=self.version,
            elements=sorted(elements, key=lambda item: item.order), metadata={"title": properties.title, "author": properties.author, "table_count": len(document.tables)})
